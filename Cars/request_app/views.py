from django.shortcuts import get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from docx.shared import RGBColor, Pt
from bs4 import BeautifulSoup
import json
import os
import re
import datetime
from django.http import JsonResponse, FileResponse
import mammoth
import pypandoc
import json
from rest_framework import viewsets
from .models import (
    CCBRequests,
    ITVerticals,
    SBUs,
    CCBScopes,
    CCBClassification,
    CCBUsers,
    CCBDecisions,
    CCBPriorities,
    CCBStatus,
    CCBSections,
    CCBRoles,
    CCBSubSections,
    CCBRequestSectionJoin,
    CCBRequestSubSectionJoin,
    CCBRoleUserJoin,
    ChangeLog,
)


from .serializers import (
    CCBRequestsSerializer,
    ITVerticalsSerializer,
    SBUsSerializer,
    CCBScopesSerializer,
    CCBClassificationSerializer,
    CCBUsersSerializer,
    CCBDecisionsSerializer,
    CCBPrioritiesSerializer,
    CCBStatusSerializer,
)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.conf import settings
import requests
from rest_framework.decorators import action
from django.http import FileResponse, Http404
import tempfile
from django.db.models import Prefetch

class CCBRequestsViewSet(viewsets.ModelViewSet):
    queryset = CCBRequests.objects.select_related(
        "it_vertical",
        "sbu",
        "scope_of_change",
        "change_class",
        "submitted_by",
        "approved_by",
        "proposal_assessed_by",
        "proposal_recommend",
        "arch_assessed_by",
        "arch_recommend",
        "cio_decision",
        "cio_priority",
        "change_status",
    ).prefetch_related(
        "sections__section",  # prefetch related sections
        "subsections__sub_section",  # prefetch related subsections
    )
    serializer_class = CCBRequestsSerializer


class RequestNameView(APIView):
    """
    Returns friendly request name for breadcrumbs.
    Example: { "name": "Purchase Order 123" }
    """

    def get(self, request, request_id):
        try:
            req = CCBRequests.objects.get(row_id=request_id)
            return Response({"name": req.change_title}, status=status.HTTP_200_OK)
        except CCBRequests.DoesNotExist:
            return Response(
                {"name": f"Request {request_id}"}, status=status.HTTP_404_NOT_FOUND
            )


class SectionSubsectionNameView(APIView):
    """
    Returns section › subsection friendly names for breadcrumbs.
    Example: { "name": "HR › Payroll" }
    """

    def get(self, request, section_id, subsection_id):

        try:
            subsection = CCBRequestSubSectionJoin.objects.get(
                row_id=subsection_id
            )
            section_name = subsection.section.section_name if subsection else f"Section {section_id}"
            subsection_name = (
            subsection.sub_section.sub_section_name if subsection else f"Subsection {subsection_id}"
        )
        except:
            section_name = f"Section {section_id}"
            subsection_name = f"Subsection {subsection_id}"

        

        # return Response(
        #     {"name": f"{section_name} › {subsection_name}"}, status=status.HTTP_200_OK
        # )
        return Response(
            {"section": section_name, "subsection": subsection_name},
            status=status.HTTP_200_OK
        )


class ITVerticalsViewSet(viewsets.ModelViewSet):
    queryset = ITVerticals.objects.all()
    serializer_class = ITVerticalsSerializer


class SBUsViewSet(viewsets.ModelViewSet):
    queryset = SBUs.objects.all()
    serializer_class = SBUsSerializer


class CCBScopesViewSet(viewsets.ModelViewSet):
    queryset = CCBScopes.objects.all()
    serializer_class = CCBScopesSerializer


class CCBClassificationViewSet(viewsets.ModelViewSet):
    queryset = CCBClassification.objects.all()
    serializer_class = CCBClassificationSerializer


class CCBUsersViewSet(viewsets.ModelViewSet):
    queryset = CCBUsers.objects.all()
    serializer_class = CCBUsersSerializer


class CCBDecisionsViewSet(viewsets.ModelViewSet):
    queryset = CCBDecisions.objects.all()
    serializer_class = CCBDecisionsSerializer


class CCBPrioritiesViewSet(viewsets.ModelViewSet):
    queryset = CCBPriorities.objects.all()
    serializer_class = CCBPrioritiesSerializer


class CCBStatusViewSet(viewsets.ModelViewSet):
    queryset = CCBStatus.objects.all()
    serializer_class = CCBStatusSerializer


class UploadDocView(APIView):

    def get(self, request, *args, **kwargs):

        path = os.path.join(settings.MEDIA_ROOT, "documents")
        file_name = "temp.docx"
        file_path = os.path.join(path, file_name)

        if not os.path.exists(file_path):
            raise Http404("File not found")

        response = FileResponse(
            open(file_path, "rb"),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename="sample.docx"'
        return response

    @staticmethod
    def upload_file_to_onedrive(local_filepath, remote_filename):
        access_token = "YOUR_MICROSOFT_GRAPH_ACCESS_TOKEN"
        url = f"https://graph.microsoft.com/v1.0/me/drive/root:/{remote_filename}:/content"
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/octet-stream",
        }
        with open(local_filepath, "rb") as f:
            data = f.read()
        response = requests.put(url, headers=headers, data=data)
        if response.status_code in (200, 201):
            return True, "Success"
        else:
            return False, response.text

    def post(self, request, format=None):
        try:
            file = request.FILES.get("file")
            if not file:
                return Response(
                    {"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST
                )

            path = os.path.join(settings.MEDIA_ROOT, "documents")
            file_name = "temp.docx"
            file_path = os.path.join(path, file_name)

            # (1) Save file locally or handle it directly
            with open(file_path, "wb+") as f:
                for chunk in file.chunks():
                    f.write(chunk)

            # (2) Upload to OneDrive
            # success, message = UploadDocView.upload_file_to_onedrive('temp.docx', file.name)
            # if success:
            #     return Response({'message': 'Uploaded successfully'}, status=status.HTTP_200_OK)
            # else:
            #     return Response({'error': message}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            return Response(
                {"message": "success"},
            )

        except Exception as e:
            return Response(
                {"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class CheckUserAPI(APIView):
    def post(self, request):
        email = request.data.get("email")
        ad_login = request.data.get("loginId")

        if not email or not ad_login:
            return Response(
                {"valid": False, "message": "Email and loginId are required"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            user = CCBUsers.objects.get(email=email, ad_login=ad_login)

            # ✅ Active roles only (both join + role active)
            roles = (
                user.roles.filter(
                    is_active=True,
                    role_user_joins__is_active=True
                ).distinct()
            )

            roles_data = []
            for role in roles:
                # ✅ Active sections only (both join + section active)
                sections = (
                    role.sections.filter(
                        is_active=True,
                        ccbsectionroles__is_active=True
                    )
                    .values("row_id", "section_name")
                )

                roles_data.append({
                    "roleId": role.row_id,
                    "roleDescription": role.role_description,
                    "sections": list(sections)
                })

            return Response(
                {
                    "valid": True,
                    "message": "User validated successfully",
                    "user": {
                        "rowId": user.row_id,
                        "email": user.email,
                        "loginId": user.ad_login,
                        "fullName": user.user_name,
                        "roles": roles_data,
                    },
                },
                status=status.HTTP_200_OK,
            )

        except CCBUsers.DoesNotExist:
            return Response(
                {"valid": False, "message": "Invalid credentials"},
                status=status.HTTP_401_UNAUTHORIZED,
            )


DOCS_DIR = os.path.join(settings.MEDIA_ROOT, "documents")
os.makedirs(DOCS_DIR, exist_ok=True)


@csrf_exempt
def import_doc(request, filename):
    """Check if DOCX exists, if not create it. Return HTML from DOCX."""
    filepath = os.path.join(DOCS_DIR, f"{filename}.docx")

    # Create file if not exists
    if not os.path.exists(filepath):
        doc = Document()
        doc.add_paragraph("New Document Created.")
        doc.save(filepath)

    # Convert DOCX -> HTML
    style_map = """
                i => em
                b => strong
                u => u
                """

    with open(filepath, "rb") as f:
        result = mammoth.convert_to_html(f, style_map=style_map)
        html_content = result.value

    return JsonResponse({"html": html_content, "filename": f"{filename}.docx"})


def rgb_to_rgbcolor(rgb_str):
    """Convert 'rgb(r, g, b)' string to RGBColor"""
    match = re.match(r"rgb\((\d+),\s*(\d+),\s*(\d+)\)", rgb_str)
    if match:
        r, g, b = map(int, match.groups())
        return RGBColor(r, g, b)
    return None


def apply_styles(run, style_str):
    """Apply inline CSS styles to a docx run"""
    if not style_str:
        return
    styles = {}
    for item in style_str.split(";"):
        if ":" in item:
            key, value = item.split(":")
            styles[key.strip()] = value.strip()

    if "color" in styles:
        color = styles["color"]
        rgb = rgb_to_rgbcolor(color)
        if rgb:
            run.font.color.rgb = rgb

    if "background-color" in styles:
        bg = styles["background-color"]
        rgb = rgb_to_rgbcolor(bg)
        if rgb:
            # Approximate with DOCX highlight (limited colors)
            run.font.highlight_color = 7  # yellow as default for highlights


def add_paragraph_with_inline_styles(doc, element):
    """
    Add a paragraph to docx with inline styles.
    Handles immediate children only to avoid duplicates.
    """
    paragraph = doc.add_paragraph()

    for node in element.children:  # immediate children only
        if isinstance(node, str):
            paragraph.add_run(node)
        elif node.name == "span":
            run = paragraph.add_run(node.get_text())
            apply_styles(run, node.get("style", ""))
        elif node.name in ["strong", "b"]:
            run = paragraph.add_run(node.get_text())
            run.bold = True
        elif node.name in ["em", "i"]:
            run = paragraph.add_run(node.get_text())
            run.italic = True
        elif node.name == "u":
            run = paragraph.add_run(node.get_text())
            run.underline = True
        elif node.name == "code":
            run = paragraph.add_run(node.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif node.name == "br":
            paragraph.add_run("\n")
        # Do NOT recurse into node.contents; handled only at immediate level

    return paragraph


# 🔹 Your OneDrive/Graph API config
TENANT_ID = (
    "25e09a3d-945e-45f8-b5d9-a84eeede55c5"  # e.g. yourcompany.onmicrosoft.com tenant ID
)
CLIENT_ID = "428e18ff-d05e-4f33-94ef-9ccdb739ea8a"  # from Azure App Registration
import os
CLIENT_SECRET = os.getenv("AZURE_SECRET")

GRAPH_SCOPE = "https://graph.microsoft.com/.default"
TOKEN_URL = f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token"
GRAPH_API_BASE = "https://graph.microsoft.com/v1.0"

ONEDRIVE_USER = "XYRE00446@coromandel.murugappa.com"


def get_access_token():
    """Get app-only access token (for OneDrive Business / SharePoint)"""
    data = {
        "client_id": CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "scope": GRAPH_SCOPE,
        "grant_type": "client_credentials",
    }
    resp = requests.post(TOKEN_URL, data=data)
    resp.raise_for_status()
    return resp.json()["access_token"]




# @csrf_exempt
# def export_doc(request):
#     """Convert HTML → DOCX and upload to OneDrive Cars_Docs folder"""
#     if request.method == "POST":
#         data = json.loads(request.body)
#         filename = data.get("filename", "output.docx")
#         html_content = data.get("html", "")

#         # Normalize HTML
#         html_content = html_content.strip() + "<p><br></p>"

#         # Convert HTML → DOCX (temporary file in OS temp folder)
#         # temp_path = os.path.join(tempfile.gettempdir(), filename)
#         temp_path = os.path.join(DOCS_DIR, filename)
#         pypandoc.convert_text(html_content, "docx", format="html", outputfile=temp_path)

#         # Add metadata
#         doc = Document(temp_path)
#         # doc.add_paragraph(f"Last edited: {datetime.datetime.now()}")
#         doc.save(temp_path)

#         # Get access token
#         access_token = get_access_token()

#         with open(temp_path, "rb") as f:
#             file_bytes = f.read()

#         headers = {
#             "Authorization": f"Bearer {access_token}",
#             "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         }

#         # 🔹 Path to Cars_Docs folder under the user’s OneDrive
#         folder_path = "Documents/Cars_Docs"
#         upload_url = f"{GRAPH_API_BASE}/users/{ONEDRIVE_USER}/drive/root:/{folder_path}/{filename}:/content"

#         resp = requests.put(upload_url, headers=headers, data=file_bytes)
#         if not resp.ok:
#             return JsonResponse(
#                 {"status": "error", "details": resp.text}, status=resp.status_code
#             )

#         file_info = resp.json()
#         return JsonResponse(
#             {
#                 "status": "success",
#                 "file": file_info["name"],
#                 "id": file_info["id"],
#                 "webUrl": file_info.get("webUrl"),
#             }
#         )


# @csrf_exempt
# def export_doc(request):
#     """Save updated HTML back to DOCX"""
#     if request.method == "POST":
#         data = json.loads(request.body)
#         filename = data.get("filename", "output.docx")
#         html_content = data.get("html", "")
#         
#         # 🟢 Normalize: ensure trailing <p><br></p>
#         html_content = html_content.strip() + "<p><br></p>"

#         filepath = os.path.join(DOCS_DIR, filename)

#         # Convert HTML -> DOCX
#         pypandoc.convert_text(html_content, "docx", format="html", outputfile=filepath)

#         # Optionally update metadata/last modified
#         doc = Document(filepath)
#         # doc.add_paragraph(f"Last edited: {datetime.datetime.now()}")
#         doc.save(filepath)

#         return JsonResponse({"status": "success", "file": filename})


@csrf_exempt
def export_doc(request):
    """Save updated HTML back to DOCX with DB-driven role/section check"""
    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "Invalid request method"}, status=405)

    try:
        data = json.loads(request.body)
        filename = data.get("filename", "output.docx")
        html_content = (data.get("html") or "").strip()
        user_id = data.get("user_id")
        section_name = data.get("sectionName")
    except (json.JSONDecodeError, TypeError):
        return JsonResponse({"status": "error", "message": "Invalid JSON"}, status=400)

    # Check if user exists
    # user = get_object_or_404(CCBUsers, row_id=user_id)

    # Check if user has an active role + section permission
    has_permission = CCBRoleUserJoin.objects.filter(
        user=user_id,          # user
        is_active=True,           # join active
        role__is_active=True,     # role active
        role__sections__is_active=True,
        role__sections__ccbsectionroles__is_active=True,
        role__sections__section_name=section_name
    ).exists()

    if not has_permission:
        return JsonResponse(
            {"status": "error", "message": "Forbidden: user cannot edit this section"},
            status=status.HTTP_403_FORBIDDEN
        )

    # Normalize HTML
    html_content += "<p><br></p>"

    filepath = os.path.join(DOCS_DIR, filename)

    # Convert HTML -> DOCX
    pypandoc.convert_text(html_content, "docx", format="html", outputfile=filepath)

    # Optional metadata
    doc = Document(filepath)
    # doc.add_paragraph(f"Last edited: {datetime.datetime.now()}")
    doc.save(filepath)

    return JsonResponse({"status": "success", "file": filename})